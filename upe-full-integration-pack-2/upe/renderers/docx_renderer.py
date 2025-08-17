
from __future__ import annotations
from docx import Document
from docx.shared import Pt
from upe.doc_ast import Doc

def render_docx(doc: Doc, out_path: str) -> None:
    d = Document()
    h1 = d.add_heading(doc.title, level=1)
    for sec in doc.sections:
        d.add_heading(sec.heading, level=2)
        for block in sec.blocks:
            if block.kind == "para":
                p = d.add_paragraph(block.text)
                p.style.font.size = Pt(12)
            elif block.kind == "list":
                for item in block.items:
                    d.add_paragraph(item, style='List Bullet')
            elif block.kind == "table":
                rows = len(block.rows)
                cols = len(block.rows[0]) if rows else 0
                if rows and cols:
                    t = d.add_table(rows=rows, cols=cols)
                    for r in range(rows):
                        for c in range(cols):
                            t.cell(r,c).text = block.rows[r][c]
    d.save(out_path)
